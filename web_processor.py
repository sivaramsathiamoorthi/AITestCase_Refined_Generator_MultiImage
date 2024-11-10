import streamlit as st
from web_test_case_generator import WebTestCaseGenerator
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document


class WebProcessor:
    def __init__(self, api_key=None):
        self.api_key = api_key or st.session_state.get("api_key")
        if not self.api_key:
            st.error("API key is required for generating responses.")
            return
        self.web_generator = WebTestCaseGenerator(api_key=self.api_key)

    def reset_session_state(self):
        """Clear session state data for a new URL input, except for API key."""
        for key in ["vector_store", "response", "related_chunks", "web_data_loaded", "response_displayed",
                    "related_displayed"]:
            st.session_state[key] = "" if key == "response" else False if key.endswith("displayed") else None

    def run(self):
        # Ensure session state variables are initialized
        st.session_state.setdefault("response", "")
        st.session_state.setdefault("related_chunks", [])
        st.session_state.setdefault("vector_store", None)
        st.session_state.setdefault("web_data_loaded", False)
        st.session_state.setdefault("response_displayed", False)
        st.session_state.setdefault("related_displayed", False)

        # Web Data Input
        url = st.text_input("Enter the website URL:")

        # Reset session state when a new URL is entered
        if url and st.button("Load Web Data"):
            self.reset_session_state()  # Clear previous data before loading new data
            message = self.web_generator.prepare_web_data(url)
            if st.session_state.get("web_data_loaded"):
                st.success("Web data loaded successfully!")
            else:
                st.error(message)

        # Question input and response generation
        if st.session_state.get("web_data_loaded"):
            question = st.text_area("Enter your question:", height=100)
            if st.button("Generate Web Response") and question:
                main_response, related_chunks = self.generate_response_with_related(question)
                st.session_state[
                    "response"] = main_response or "No response generated. Please check your question and try again."
                st.session_state[
                    "related_chunks"] = related_chunks or []  # Ensure related_chunks is an empty list if None
                st.session_state["response_displayed"] = True
                st.session_state["related_displayed"] = True

        # Display the generated response and related information if available
        if st.session_state.get("response_displayed"):
            st.markdown("### Generated Response:")
            formatted_response = self.format_response(st.session_state.response)
            st.markdown(formatted_response, unsafe_allow_html=True)

            # Copy to Clipboard Button
            if st.button("Copy to Clipboard"):
                st.write(
                    f"<script>navigator.clipboard.writeText(`{st.session_state['response']}`);</script>",
                    unsafe_allow_html=True
                )
                st.success("Response copied to clipboard!")

        # Display related information if available
        if st.session_state.get("related_displayed") and st.session_state.related_chunks:
            self.display_related_information(st.session_state.related_chunks)

        # Download options for PDF and Word document
        col1, col2 = st.columns(2)
        with col1:
            pdf_data = self.generate_pdf_with_related(st.session_state.response, st.session_state.related_chunks)
            st.download_button(
                label="Download as PDF",
                data=pdf_data,
                file_name="response.pdf",
                mime="application/pdf"
            )
        with col2:
            docx_data = self.generate_word_doc_with_related(st.session_state.response, st.session_state.related_chunks)
            st.download_button(
                label="Download as Word Document",
                data=docx_data,
                file_name="response.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    def generate_response_with_related(self, question):
        """Generates the main response and retrieves related information chunks with images if available."""
        main_response = self.web_generator.generate_response(question)

        # Fetch related information chunks with specified search_type
        related_chunks = []
        documents = st.session_state.vector_store.search(question, k=3, search_type="similarity")

        for doc in documents:
            title = doc.metadata.get("title", "Related Information")
            summary = " ".join(doc.page_content.split()[:40])  # Simple summary, taking first 40 words
            image_url = doc.metadata.get("image_url")  # Assuming metadata might contain 'image_url'

            chunk = {
                "title": title,
                "summary": summary,
                "image_url": image_url  # Include image URL if available
            }
            related_chunks.append(chunk)

        return main_response, related_chunks

    def display_related_information(self, related_chunks):
        """Displays related information with optional images."""
        st.markdown("### Related Information:")
        for chunk in related_chunks:
            st.markdown(f"**{chunk['title']}**")
            st.markdown(chunk["summary"])
            if chunk["image_url"]:
                st.image(chunk["image_url"], use_column_width=True)

    def format_response(self, text):
        formatted_text = ""
        lines = text.splitlines()
        for line in lines:
            if line.startswith("1. ") or line.startswith("2. "):
                formatted_text += f"<li>{line[3:]}</li>"
            elif "**" in line:
                formatted_text += f"<p><b>{line}</b></p>"
            else:
                formatted_text += f"<p>{line}</p>"

        return f"""
            <div style="background-color: #f0f4f8; padding: 20px; border-radius: 8px; border: 1px solid #ddd;
                        font-family: Arial, sans-serif; color: #333; max-height: 300px; overflow-y: scroll;">
                {formatted_text}
            </div>
        """

    def generate_pdf_with_related(self, text, related_chunks):
        """Generates a PDF with the main response and related information."""
        pdf_buffer = BytesIO()
        pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
        pdf.setFont("Helvetica", 12)

        width, height = letter
        y_position = height - 40

        # Write main response
        for line in text.splitlines():
            if y_position <= 40:
                pdf.showPage()
                y_position = height - 40
            pdf.drawString(40, y_position, line)
            y_position -= 14

        # Write related chunks
        for chunk in related_chunks or []:  # Handle potential None by defaulting to an empty list
            pdf.drawString(40, y_position, f"Title: {chunk['title']}")
            y_position -= 14
            pdf.drawString(40, y_position, f"Summary: {chunk['summary']}")
            y_position -= 14
            if y_position <= 40:
                pdf.showPage()
                y_position = height - 40

        pdf.save()
        pdf_buffer.seek(0)
        return pdf_buffer

    def generate_word_doc_with_related(self, text, related_chunks):
        """Generates a Word document with the main response and related information."""
        doc_buffer = BytesIO()
        doc = Document()
        doc.add_paragraph("Generated Response:")
        doc.add_paragraph(text)

        doc.add_paragraph("Related Information:")
        for chunk in related_chunks or []:  # Handle potential None by defaulting to an empty list
            doc.add_paragraph(f"Title: {chunk['title']}")
            doc.add_paragraph(f"Summary: {chunk['summary']}")

        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer
